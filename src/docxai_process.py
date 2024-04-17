import os
import time
import logging
import json
import requests

from coolname import generate_slug
from pdf2docx import Converter

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from .save_formatting import copy_style, apply_style
from .utilites import load_env


env_vars = load_env()

YANDEX_PRE_PROMPT = """\n
Ты — ИИ-помощник, тебе на вход приходит JSON файл, ты должен изменить его в соответствии со следующими правилами и инструкциями:
0. Проанализировать загруженный текст формата json. В тексте каждого элемента json при необходимости выполнить следующие действия (не меняя структуру json файла).
"""

YANDEX_POST_PROMPT = '''\n
В конце - выведи изменённый текст так же в формате json, сохранив всю структуру, количество элементов и нумерацию.
Дополнение к инструкциям. Очен важно:
- Выведи изменённый текст так же в формате json, сохранив всю структуру, количество элементов и нумерацию.
- Не удаляй лишнее из оригинала. Изменениям должны подвергнуться лишь неопределенности. Все четко-определенные элементы должны остаться в том же виде в котором были.
- Не удаляй все НЕ буквенные символы, даже если элемент json целиком состоит из них. Например: "10": "__________" должна так и остаться "10": "__________".
- Обязательно в финальном тексте должен быть абсолютно весь текст из оригинального тексте с учётом корректировок и ничего дополнительно писать не надо.
- Ответ нужно предоставить в JSON формате без каких либо комментариев вначале и без каких либо комментариев в конце. Только JSON, начинается с "{" заканчивается "}".
'''

YANDEXGPT_TOKEN = env_vars["yandexgpt_token"] 
FOLDER_ID = env_vars["folder_id"]

def pdf_converter(pdf_path) -> str:
    docx_path = pdf_path.replace(f".pdf", f"-{generate_slug(2)}.docx")
    try:
        c = Converter(pdf_path)
        c.convert(docx_path)
        c.close()
    except Exception as e: # TODO
        print("")
        return ""
    else:
        return docx_path

def file_process(file_path) -> str:
    if file_path.endswith(".docx"):
        new_file_path = file_path.replace(f".docx", f"-{generate_slug(2)}.docx")
        logging.info(f"Document process start: {file_path} -> {new_file_path}")
        f = open(file_path, 'rb')
        doc = Document(f)
        f.close()
        new_file_path = main_process_ai(doc, new_file_path)
        return new_file_path
    elif file_path.endswith(".pdf"):
        new_file_path = pdf_converter(file_path)
        logging.info(f"Document process start: {file_path} -> {new_file_path}")
        logging.error(f"Not supported format: {file_path}")
        ...
        return ''
    else:
        new_file_path = ''
        logging.error(f"Not supportet format: {file_path}")
        return new_file_path

def make_paragraphs_dict(doc_paragraphs:list)->dict:
    json_dicts = dict()
    for i in range(len(doc_paragraphs)):
        json_dicts[i] = doc_paragraphs[i].text
    return json_dicts

def make_tables_dict(doc)->dict:
    table_dict = dict()
    for i in range(len(doc.tables)):
        table_dict[i] = {}
        for j in range(len(doc.tables[i].table._cells)):
            table_dict[i][j] = make_paragraphs_dict(doc.tables[i].table._cells[j].paragraphs)
    return table_dict

def write_changes_paragraph(doc, paragrhaph_dict_modified:dict):
    for i in range(len(doc.paragraphs)):
        if not paragrhaph_dict_modified.get(str(i)) or (paragrhaph_dict_modified[str(i)] == doc.paragraphs[i].text):
            continue
        else:
            runs_amount = len(doc.paragraphs[i].runs)
            run_number = -1  #careful
            mod_flag = True

            tmp_original_text = doc.paragraphs[i].text
            for n in range(len(doc.paragraphs[i].runs)):
                doc.paragraphs[i].runs[n].text = ''
            style = copy_style(doc.paragraphs[i].runs[0])
            logging.info(f"original text: {tmp_original_text}")
            logging.info(f"modified text: {paragrhaph_dict_modified[str(i)]}")
            for word in find_changes_generator(tmp_original_text, paragrhaph_dict_modified[str(i)]):
                if word.startswith("###"):
                    if not mod_flag or run_number == -1:
                        run_number+=1
                        mod_flag = True
                    if run_number >= runs_amount:
                        doc.paragraphs[i].add_run()
                        apply_style(doc.paragraphs[i].runs[run_number], style)
                        doc.paragraphs[i].runs[run_number].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    doc.paragraphs[i].runs[run_number].text += f"{word} "[3:]
                    doc.paragraphs[i].runs[run_number].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

                else:
                    if mod_flag or run_number == -1:
                        run_number+=1
                        mod_flag = False
                    if run_number >= runs_amount:
                        doc.paragraphs[i].add_run()
                        apply_style(doc.paragraphs[i].runs[run_number], style)
                    doc.paragraphs[i].runs[run_number].text += f"{word} "

def write_changes_table(doc, table_dict_modified:dict):
    for i in range(len(doc.tables)):
        for j in range(len(doc.tables[i].table._cells)):
            for n in range(len(doc.tables[i].table._cells[j].paragraphs)):
                if not table_dict_modified[i].get(str(j)) or not table_dict_modified[i][str(j)].get(str(n)) or (table_dict_modified[i][str(j)][str(n)] == doc.tables[i].table._cells[j].paragraphs[n].text):
                    continue
                else:
                    runs_amount = len(doc.tables[i].table._cells[j].paragraphs[n].runs)
                    run_number = -1  #careful
                    mod_flag = True

                    tmp_original_text = doc.tables[i].table._cells[j].paragraphs[n].text
                    for r in range(len(doc.tables[i].table._cells[j].paragraphs[n].runs)):
                        doc.tables[i].table._cells[j].paragraphs[n].runs[r].text = ''
                    style = copy_style(doc.tables[i].table._cells[j].paragraphs[n].runs[0])
                    logging.info(f"original text: {tmp_original_text}")
                    logging.info(f"modified text: {table_dict_modified[i][str(j)][str(n)]}")
                    for word in find_changes_generator(str(tmp_original_text), str(table_dict_modified[i][str(j)][str(n)])):
                        if word.startswith("###"):
                            if not mod_flag or run_number == -1:
                                run_number+=1
                                mod_flag = True
                            if run_number >= runs_amount:
                                doc.tables[i].table._cells[j].paragraphs[n].add_run()
                                apply_style(doc.tables[i].table._cells[j].paragraphs[n].runs[run_number], style)
                                doc.tables[i].table._cells[j].paragraphs[n].runs[run_number].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                            doc.tables[i].table._cells[j].paragraphs[n].runs[run_number].text += f"{word} "[3:]
                            doc.tables[i].table._cells[j].paragraphs[n].runs[run_number].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

                        else:
                            if mod_flag or run_number == -1:
                                run_number+=1
                                mod_flag = False
                            if run_number >= runs_amount:
                                doc.tables[i].table._cells[j].paragraphs[n].add_run()
                                apply_style(doc.tables[i].table._cells[j].paragraphs[n].runs[run_number], style)
                            doc.tables[i].table._cells[j].paragraphs[n].runs[run_number].text += f"{word} " 

def main_process_ai(doc, new_file_path_name):
    paragrhaph_dict = make_paragraphs_dict(doc.paragraphs)
    paragrhaph_dict_modified = func_yandexgpt(paragrhaph_dict)
    table_dict = make_tables_dict(doc)
    table_dict_modified = yandex_gpt_tables(table_dict)
    write_changes_paragraph(doc, paragrhaph_dict_modified)
    write_changes_table(doc, table_dict_modified) 
    
    doc.save(new_file_path_name)
    return new_file_path_name

def find_changes_generator(text_1, text_2):
    tokens_1 = text_1.split(" ")
    tokens_2 = text_2.split(" ")
    length_1 = len(tokens_1)
    length_2 = len(tokens_2)
    index_or = 0
    index_mod = 0
    iteration = 0
    flag = True
    while(index_mod != length_2):
        if (tokens_1[index_or] != tokens_2[index_mod]):
            for tmp_index in range(index_or, length_1):
                if tokens_1[tmp_index] == tokens_2[index_mod]:
                    index_or = tmp_index
                    flag = False
                    if iteration == 0:
                        yield "###" + tokens_2[index_mod]
                    else:
                        yield tokens_2[index_mod]
                    break
            if flag:
                yield "###" + tokens_2[index_mod]
                index_mod+=1
                iteration += 1
        
        else:
            if flag:
                yield tokens_2[index_mod]
            index_mod+=1
            if index_or < length_1 - 1:
                index_or+=1
            flag = True
            iteration = 0

def get_prompt() -> str:
    try:
        with open("prompt.txt", "r") as file:
            prompt = file.read()
            return prompt
    except Exception as e:
        logging.error(f"read prompt fail: {e}")

def get_instruction() -> str:
    try:
        with open("instruction.txt", "r") as file:
            instruction = file.read()
            return instruction
    except Exception as e:
        logging.error(f"read instruction fail: {e}")

def send_prompt(chunk:str, PROMPT:str, INSTRUCTION:str, token:str, folder_id:str) -> str:
    success = False
    response = None
    logging.info(f"starting work with a chunk: {chunk}")
    time.sleep(2)
    user_message = {
                    "modelUri": f"gpt://b1gr06ae9rrolg6nf1c4/yandexgpt-pro",
                    "completionOptions": {
                        "stream": False,
                        "temperature": 0.6,
                        "maxTokens": "1000"
                    },
                    "messages": [
                        {
                        "role": "system",
                        "text": INSTRUCTION
                        },
                        {
                        "role": "assistant",
                        "text": PROMPT
                        },
                        {
                        "role": "user",
                        "text": chunk
                        }
                    ]
                    }
    
    for _ in range(4):
        time.sleep(2)
        response = requests.post(
                                url     = f"https://llm.api.cloud.yandex.net/foundationModels/v1/completion",
                                headers = {'Content-Type': 'application/json', 
                                    'Authorization': 'Api-Key ' + token, 
                                    'x-folder-id': folder_id},
                                json = user_message
                            )
        
        if response.status_code == 200:
            logging.info(f"response status code: {response.status_code}")
        else:
            logging.warning(f"YandexGPT problem, status code: {response.status_code}")
            continue
        try:
            tmp_json_response = json.loads(response.text)
            tmp_json = json.loads(tmp_json_response['result']['alternatives'][0]['message']['text'].strip('[]`json\n '))
            tmp_orig_dict = json.loads(chunk)
            for key in tmp_orig_dict.keys():
                tmp_json[key]
            success = True
            break
        except:
            logging.warning("JSON is not valid or has missing elements, new try")
            response = requests.post(
                        url     = f"https://llm.api.cloud.yandex.net/foundationModels/v1/completion",
                        headers = {'Content-Type': 'application/json', 
                            'Authorization': 'Api-Key ' + token, 
                            'x-folder-id': folder_id},
                        json = user_message
                )
    
    if response.status_code != 200:
            logging.error(f"this chunck wiil be missed due to Yandexgpt problems: {chunk}")
    if not success:
        logging.error(f"this chunck wiil be missed due to incorrect JSON formatting or missing elements: {chunk}")
    return response.text

def func_yandexgpt(input_dict:dict)->dict:
    PROMPT = YANDEX_PRE_PROMPT + get_prompt() + YANDEX_POST_PROMPT
    INSTRUCTION = get_instruction()

    token = YANDEXGPT_TOKEN
    folder_id = FOLDER_ID
    chunk_size=7
    raw_str_json = ''
    json_strings = []
    items = list(input_dict.items())
    
    for i in range(0, len(items), chunk_size):
        chunk_dict = dict(items[i:i+chunk_size])
        json_string = json.dumps(chunk_dict, ensure_ascii=False)
        json_strings.append(json_string)
    
    for chunk in json_strings:
        response = send_prompt(chunk, PROMPT, INSTRUCTION, token, folder_id)
        if response:
            response_json = json.loads(response)
            logging.info(f"json text {response_json['result']['alternatives'][0]['message']['text']}")
            raw_str_json += (response_json['result']['alternatives'][0]['message']['text'].strip('[]}{`json\n ')) + ','

    final_json_str = "{"+ raw_str_json.strip(",") + "}"
    logging.info(f"FILAL JSON str: {final_json_str}")
    final_json = json.loads(final_json_str)
    return final_json

def yandex_gpt_tables(input_dict:dict)->dict:
    PROMPT = YANDEX_PRE_PROMPT + get_prompt() + YANDEX_POST_PROMPT
    INSTRUCTION = get_instruction()

    token = YANDEXGPT_TOKEN
    folder_id = FOLDER_ID
    chunk_size=7
    result_dict = dict()
    for main_i in input_dict.keys():
        raw_str_json = ''
        json_strings = []
        items = list(input_dict[main_i].items())
        
        for i in range(0, len(items), chunk_size):
            chunk_dict = dict(items[i:i+chunk_size])
            json_string = json.dumps(chunk_dict, ensure_ascii=False)
            json_strings.append(json_string)
        
        for chunk in json_strings:
            response = send_prompt(chunk, PROMPT, INSTRUCTION, token, folder_id)
            if response:
                response_json = json.loads(response)
                logging.info(f"json text {response_json['result']['alternatives'][0]['message']['text']}")
                raw_str_json += (response_json['result']['alternatives'][0]['message']['text'].strip('[]}{`json\n ')) + '}' + ','

        final_json_str = "{"+ raw_str_json.strip(",") + "}"
        logging.info(f"FILAL JSON str: {final_json_str}")
        final_json = json.loads(final_json_str)
        result_dict[main_i] = final_json
    return result_dict

def init_instruction():
    try:
        with open("docxai/default_prompt.txt", "r") as file:
            def_prompt = file.read()
    except Exception as e:
        logging.error(f"read prompt fail: {e}")
    try:
        with open("prompt.txt", "w") as file:
            file.write(def_prompt)
    except Exception as e:
        logging.error(f"write prompt fail: {e}")
    try:
        with open("docxai/default_instruction.txt", "r") as file:
            def_instruction = file.read()
    except Exception as e:
        logging.error(f"read instruction fail: {e}")
    try:
        with open("instruction.txt", "w") as file:
            file.write(def_instruction)
    except Exception as e:
        logging.error(f"write instruction fail: {e}")

def main():
    init_instruction() 
    file_process("test.docx")

if __name__ == '__main__':
    main()